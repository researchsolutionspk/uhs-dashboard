#!/usr/bin/env python3
"""
UHS Dashboard Generator
=======================
Run this script each time new survey data is added to the CSV files.

Usage:
    python generate.py           # regenerate + auto-push to GitHub
    python generate.py --local   # regenerate only (no git push)

Reads:
    Data/Boys Data/Boy's Questionnaire - Endline Empowering Girls_WIDE.csv
    Data/Girls Data/Girl's Questionnaire - Endline Empowering Girls_WIDE.csv

Writes:
    index.html  (regenerated from template.html with fresh data)
"""

import csv
import sys
import subprocess
from datetime import datetime
from pathlib import Path

# ── Paths ──────────────────────────────────────────────────────────────────
BASE       = Path(__file__).parent
BOYS_CSV   = BASE / "Data/Boys Data/Boy's Questionnaire - Endline Empowering Girls_WIDE.csv"
GIRLS_CSV  = BASE / "Data/Girls Data/Girl's Questionnaire - Endline Empowering Girls_WIDE.csv"
TEMPLATE   = BASE / "template.html"
OUTPUT     = BASE / "index.html"

AUTO_PUSH  = "--local" not in sys.argv

# ── Helpers ────────────────────────────────────────────────────────────────
def read_csv(path):
    with open(path, "r", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))

def sf(val, default=None):
    """Safe float — returns None for blank/zero."""
    try:
        v = float(val)
        return v if v > 0 else default
    except Exception:
        return default

def mean_of(rows, key):
    vals = [sf(r.get(key, "")) for r in rows]
    vals = [v for v in vals if v is not None]
    return round(sum(vals) / len(vals), 2) if vals else 0

def count_val(rows, key, val):
    return sum(1 for r in rows if r.get(key, "") == str(val))

def count_vals(rows, key, vals):
    return [count_val(rows, key, v) for v in vals]

def grade_band(rows, key, bins):
    """Bin a float column into ranges defined by (low, high) tuples."""
    vals = [sf(r.get(key, "")) for r in rows]
    vals = [v for v in vals if v is not None]
    return [sum(1 for v in vals if lo <= v < hi) for lo, hi in bins]

def pct(count, total):
    return round(count / total * 100) if total else 0

def js_list(lst):
    return "[" + ", ".join(str(round(v, 2)) for v in lst) + "]"

# ── Load data ──────────────────────────────────────────────────────────────
print("Reading CSV files...")
boys  = read_csv(BOYS_CSV)
girls = read_csv(GIRLS_CSV)

def _school_ids(rows):
    return {r.get('school_id', '').strip() for r in rows if r.get('school_id', '').strip()}
n_schools = len(_school_ids(boys) | _school_ids(girls))
n_boys  = len(boys)
n_girls = len(girls)
n_total = n_boys + n_girls
print(f"  Boys : {n_boys} records")
print(f"  Girls: {n_girls} records")
print(f"  Total: {n_total} records")

# ── DEMOGRAPHICS ───────────────────────────────────────────────────────────
boys_age_mean  = mean_of(boys,  "age")
girls_age_mean = mean_of(girls, "age")
boys_sib_mean  = mean_of(boys,  "siblings")
girls_sib_mean = mean_of(girls, "siblings")

def _minmax(rows, key):
    vals = [sf(r.get(key, "")) for r in rows]
    vals = [v for v in vals if v is not None]
    return (int(min(vals)), int(max(vals))) if vals else (0, 0)
boys_age_min,  boys_age_max  = _minmax(boys,  "age")
girls_age_min, girls_age_max = _minmax(girls, "age")

# Home ownership — asset_7 = House/apartment
boys_home_own_pct  = pct(count_val(boys,  "asset_7", 1), n_boys)
girls_home_own_pct = pct(count_val(girls, "asset_7", 1), n_girls)

# ── ACADEMIC ───────────────────────────────────────────────────────────────
bg = [sf(r.get("percentage_grade7","")) for r in boys  if sf(r.get("percentage_grade7","")) is not None]
gg = [sf(r.get("percentage_grade7","")) for r in girls if sf(r.get("percentage_grade7","")) is not None]

boys_grade_mean  = round(sum(bg)/len(bg), 1) if bg else 0
girls_grade_mean = round(sum(gg)/len(gg), 1) if gg else 0
boys_grade_min   = int(min(bg)) if bg else 0
girls_grade_min  = int(min(gg)) if gg else 0
boys_grade_max   = round(max(bg), 1) if bg else 0
girls_grade_max  = round(max(gg), 1) if gg else 0

grade_bins = [(30,40),(40,50),(50,60),(60,70),(70,80),(80,101)]
boys_grade_dist  = grade_band(boys,  "percentage_grade7", grade_bins)
girls_grade_dist = grade_band(girls, "percentage_grade7", grade_bins)

boys_perf_bands  = grade_band(boys,  "percentage_grade7", [(0,45),(45,60),(60,75),(75,101)])
girls_perf_bands = grade_band(girls, "percentage_grade7", [(0,45),(45,60),(60,75),(75,101)])

ba_vals = [sf(r.get("absence_days","")) for r in boys  if sf(r.get("absence_days","")) is not None]
ga_vals = [sf(r.get("absence_days","")) for r in girls if sf(r.get("absence_days","")) is not None]
boys_abs_mean  = round(sum(ba_vals)/len(ba_vals),1) if ba_vals else 0
girls_abs_mean = round(sum(ga_vals)/len(ga_vals),1) if ga_vals else 0

# Aspirations [matric=3, inter=4, ba=5, ma=6, pro=7]
boys_asp  = count_vals(boys,  "highest_edu", [3,4,5,6,7])
girls_asp = count_vals(girls, "highest_edu", [3,4,5,6,7])
girls_asp_pro_pct = pct(count_val(girls, "highest_edu", 7), n_girls)
boys_asp_pro_pct  = pct(count_val(boys,  "highest_edu", 7), n_boys)
grade_gap = round(girls_grade_mean - boys_grade_mean, 1)

# ── SAFETY ─────────────────────────────────────────────────────────────────
boys_commute_mean  = mean_of(boys,  "safe_commuting_school")
girls_commute_mean = mean_of(girls, "safe_commuting_school")
boys_school_mean   = mean_of(boys,  "safe_school_neighbourhood")
girls_school_mean  = mean_of(girls, "safe_school_neighbourhood")
boys_home_mean     = mean_of(boys,  "safe_home_neighbourhood")
girls_home_mean    = mean_of(girls, "safe_home_neighbourhood")

boys_commute_dist  = count_vals(boys,  "safe_commuting_school",     [1,2,3,4])
girls_commute_dist = count_vals(girls, "safe_commuting_school",     [1,2,3,4])
boys_school_dist   = count_vals(boys,  "safe_school_neighbourhood", [1,2,3,4])
girls_school_dist  = count_vals(girls, "safe_school_neighbourhood", [1,2,3,4])
boys_home_dist     = count_vals(boys,  "safe_home_neighbourhood",   [1,2,3,4])
girls_home_dist    = count_vals(girls, "safe_home_neighbourhood",   [1,2,3,4])

# ── BEHAVIOUR MODIFICATIONS (girls) ────────────────────────────────────────
girls_route    = count_vals(girls, "route_changed",    [1,2,3,4])
girls_timing   = count_vals(girls, "time_changed",     [1,2,3,4])
girls_clothing = count_vals(girls, "clothing_changed", [1,2,3,4])
girls_escape   = count_vals(girls, "identify_escape",  [1,2,3,4])
girls_clothing_mod_pct = pct(girls_clothing[0]+girls_clothing[1], n_girls)
girls_route_mod_pct    = pct(girls_route[0]+girls_route[1],       n_girls)
girls_timing_mod_pct   = pct(girls_timing[0]+girls_timing[1],     n_girls)
girls_clothing_mod_n   = girls_clothing[0] + girls_clothing[1]
girls_escape_pct       = pct(girls_escape[0]+girls_escape[1],     n_girls)

# ── MOBILITY ───────────────────────────────────────────────────────────────
boys_transport    = count_vals(boys,  "school_transport", [1,2,3])
girls_transport   = count_vals(girls, "school_transport", [1,2])
boys_travel_mean  = mean_of(boys,  "school_time")
girls_travel_mean = mean_of(girls, "school_time")

# ── FAMILY BACKGROUND ──────────────────────────────────────────────────────
boys_father_edu  = count_vals(boys,  "father_edu", [1,2,3,4,5,6])
girls_father_edu = count_vals(girls, "father_edu", [1,2,3,4,5,6])
boys_mother_edu  = count_vals(boys,  "mother_edu", [1,2,3,4,6])
girls_mother_edu = count_vals(girls, "mother_edu", [1,2,3,4,6])

boys_father_occ  = count_vals(boys,  "occupation_father", [1,2,3,4,10,88])
girls_father_occ = count_vals(girls, "occupation_father", [1,2,3,4,10,88])
boys_mother_occ  = count_vals(boys,  "occupation_mother", [1,2,3,8,99])
girls_mother_occ = count_vals(girls, "occupation_mother", [1,2,3,8,99])

# Father education — Intermediate = code 5 (modal); codes 1,2 = no/primary only
boys_father_inter_pct  = pct(count_val(boys,  "father_edu", 5), n_boys)
girls_father_inter_pct = pct(count_val(girls, "father_edu", 5), n_girls)
boys_father_lowed_pct  = pct(count_val(boys,  "father_edu", 1) + count_val(boys,  "father_edu", 2), n_boys)
girls_father_lowed_pct = pct(count_val(girls, "father_edu", 1) + count_val(girls, "father_edu", 2), n_girls)

# Mother homemaker — occupation_mother code 8
boys_mother_home_pct  = pct(count_val(boys,  "occupation_mother", 8), n_boys)
girls_mother_home_pct = pct(count_val(girls, "occupation_mother", 8), n_girls)

# ── ASSETS ─────────────────────────────────────────────────────────────────
asset_vars = [f"asset_{i}" for i in [1,2,3,4,5,6,7,8,9,11,12,13,14]]
boys_assets_pct  = [pct(count_val(boys,  a, 1), n_boys)  for a in asset_vars]
girls_assets_pct = [pct(count_val(girls, a, 1), n_girls) for a in asset_vars]

# ── LANGUAGE ───────────────────────────────────────────────────────────────
bl = count_vals(boys, "language", [1,2,7,99])
boys_lang = [bl[0], bl[1], bl[2]+bl[3]]   # Urdu, Punjabi, Other

# ── PQL WELLBEING ──────────────────────────────────────────────────────────
pql_vars = ["pql_2","pql_3","pql_5","pql_6","pql_15"]
boys_pql_means  = [mean_of(boys,  v) for v in pql_vars]
girls_pql_means = [mean_of(girls, v) for v in pql_vars]

def dist5(rows, var):
    return count_vals(rows, var, [1,2,3,4,5])

boys_pql_dists  = [dist5(boys,  v) for v in pql_vars]
girls_pql_dists = [dist5(girls, v) for v in pql_vars]

boys_pql_overall  = round(sum(boys_pql_means) /len(boys_pql_means),  2)
girls_pql_overall = round(sum(girls_pql_means)/len(girls_pql_means), 2)

# ── SELF-EFFICACY ──────────────────────────────────────────────────────────
se_vars = ["self_efficacy_1","self_efficacy_2","self_efficacy_3"]
boys_se_means  = [mean_of(boys,  v) for v in se_vars]
girls_se_means = [mean_of(girls, v) for v in se_vars]

# ── GENDER NORMS ───────────────────────────────────────────────────────────
gn_vars = ["gender_2","gender_3","gender_5","gender_6"]
boys_gn_means  = [mean_of(boys,  v) for v in gn_vars]
girls_gn_means = [mean_of(girls, v) for v in gn_vars]
boys_gn_dists  = [dist5(boys,  v) for v in gn_vars]
girls_gn_dists = [dist5(girls, v) for v in gn_vars]

# ── HARASSMENT ─────────────────────────────────────────────────────────────
boys_phys_witness  = count_val(boys,  "physical_other_1", 1)
girls_phys_witness = count_val(girls, "physical_other_1", 1)
boys_verbal_witness  = count_val(boys,  "verbal_other_1", 1)
girls_verbal_witness = count_val(girls, "verbal_other_1", 1)
boys_verbal_pct  = pct(boys_verbal_witness,  n_boys)
girls_verbal_pct = pct(girls_verbal_witness, n_girls)

boys_blame  = count_vals(boys,  "touch_fault", [1,2,3,4,5])
girls_blame = count_vals(girls, "touch_fault", [1,2,3,4,5])
# Strong perpetrator attribution = code 5 (strongly agree perpetrator at fault)
boys_blame_perp_pct  = pct(boys_blame[4],  n_boys)
girls_blame_perp_pct = pct(girls_blame[4], n_girls)

# ── MASCULINITY NORMS ──────────────────────────────────────────────────────
masc_vars = [f"masculinity_{i}" for i in range(1,6)]
masc_means = [mean_of(boys, v) for v in masc_vars]
masc_dists = [dist5(boys, v) for v in masc_vars]
masc_pcts  = [round(m/5*100, 1) for m in masc_means]

# ── IMPULSE CONTROL ────────────────────────────────────────────────────────
impulse_means = [mean_of(boys, f"impulse_{i}") for i in range(1,9)]

# ── CBT & PEER PRESSURE ────────────────────────────────────────────────────
cbt_means  = [mean_of(boys, f"cbt_{i}") for i in range(1,4)]
peer_means = [mean_of(boys, f"peer_press_{i}") for i in range(1,4)]

# ── GIRLS TIME USE ─────────────────────────────────────────────────────────
time_vars = ["time_water","time_cooking","time_cleaning","time_caring","time_business"]
girls_time_use = [mean_of(girls, v) for v in time_vars]

# ── HARASSMENT SHARING ─────────────────────────────────────────────────────
girls_share_nobody = count_val(girls, "harassment_share_1", 1)
girls_share_friend = count_val(girls, "harassment_share_3", 1)
girls_share_parent = count_val(girls, "harassment_share_2", 1)
girls_report       = count_val(girls, "harassment_report",  1)
girls_report_pct   = pct(girls_report, n_girls)
girls_share_nobody_pct = pct(girls_share_nobody, n_girls)
girls_share_friend_pct = pct(girls_share_friend, n_girls)
girls_share_parent_pct = pct(girls_share_parent, n_girls)

# ── RADAR SCORES (0-100) ───────────────────────────────────────────────────
boys_radar = [
    boys_grade_mean,
    round((boys_commute_mean - 1) / 3 * 100, 1),
    round((boys_pql_overall  - 1) / 4 * 100, 1),
    round((sum(boys_se_means)/3  - 1) / 2 * 100, 1),
    round((1 - (boys_gn_means[0] - 1) / 4) * 100, 1),
]
girls_radar = [
    girls_grade_mean,
    round((girls_commute_mean - 1) / 3 * 100, 1),
    round((girls_pql_overall  - 1) / 4 * 100, 1),
    round((sum(girls_se_means)/3  - 1) / 2 * 100, 1),
    round((1 - (girls_gn_means[0] - 1) / 4) * 100, 1),
]

# ── TIMESTAMP ──────────────────────────────────────────────────────────────
last_updated = datetime.now().strftime("%d %B %Y, %H:%M")

# ══════════════════════════════════════════════════════════════════════════
# SAMPLE COVERAGE  (new tab — classify endline surveys against per-school rosters)
# ══════════════════════════════════════════════════════════════════════════
# Data sources:
#   - Sample/UHS Schools with GeoCoordinates.xlsx  -> treatment/control mapping
#   - Sample/Sample Tables/*.docx                   -> per-school Table 1 / Table 2 rosters
#
# Per user: only count buckets for schools where a docx roster file exists.
# Numbers update dynamically as new docx files are dropped into Sample/Sample Tables.
print("Computing sample coverage...")

from collections import defaultdict

SAMPLE_XLSX = BASE / "Sample/UHS Schools with GeoCoordinates.xlsx"
SAMPLE_DOCX_DIR = BASE / "Sample/Sample Tables"

def _norm_id(v):
    """Normalize an ID value (from CSV or docx) to a canonical string key."""
    if v is None: return ""
    s = str(v).strip()
    if not s or s.lower() in ("nan", "none", "null"): return ""
    # Strip trailing ".0" from numeric-float strings
    if s.endswith(".0") and s[:-2].isdigit(): s = s[:-2]
    return s

def _load_school_mapping():
    """Return list of dicts: {school_id, name, emis, trt, type, category}."""
    if not SAMPLE_XLSX.exists():
        print(f"  WARNING: {SAMPLE_XLSX.name} not found — sample coverage will be empty.")
        return []
    import openpyxl
    wb = openpyxl.load_workbook(SAMPLE_XLSX, data_only=True)
    ws = wb.active
    headers = [(c.value or "").strip() if isinstance(c.value, str) else str(c.value or "").strip() for c in ws[1]]
    out = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        d = dict(zip(headers, row))
        sid = d.get("School ID")
        if sid is None: continue
        try: sid = int(sid)
        except: continue
        type_ = str(d.get("SchoolType", "")).strip()
        trt   = str(d.get("Treatment", "")).strip()
        if type_ == "Female" and trt in ("T_Girls", "T_Boys"): cat = "girls_treatment"
        elif type_ == "Female" and trt == "C":                 cat = "girls_control"
        elif type_ == "Male"   and trt == "T_Boys":            cat = "boys_treatment"
        elif type_ == "Male"   and trt == "C":                 cat = "boys_control"
        else: cat = None
        out.append({
            "school_id": sid,
            "name": str(d.get("SchoolName", "")).strip().upper(),
            "emis": str(d.get("EMIS Code", "")).strip(),
            "trt":  trt, "type": type_, "category": cat,
        })
    return out

def _docx_school_name(doc):
    """Extract school name from 'School Name:' cell in Table 1 header."""
    if not doc.tables: return ""
    for row in doc.tables[0].rows[:3]:
        for cell in row.cells:
            t = cell.text.strip()
            if "school name" in t.lower():
                cleaned = t.split(":", 1)[-1] if ":" in t else t
                return cleaned.replace("\n", " ").strip().upper()
    return ""

def _docx_table_student_ids(table):
    """Extract the set of Student IDs from a docx table (skip header and empty cells)."""
    id_col = header_row = None
    for r_idx, row in enumerate(table.rows[:3]):
        for c_idx, cell in enumerate(row.cells):
            if "student id" in cell.text.strip().lower():
                id_col = c_idx; header_row = r_idx; break
        if id_col is not None: break
    if id_col is None: return set()
    ids = set()
    for row in table.rows[header_row + 1:]:
        if id_col >= len(row.cells): continue
        val = _norm_id(row.cells[id_col].text)
        if val: ids.add(val)
    return ids

def _build_rosters(mapping):
    """Parse every .docx in SAMPLE_DOCX_DIR, fuzzy-match by school name to mapping.
    Returns {school_id: {'t1': set, 't2': set, 'files': [...]}}."""
    if not SAMPLE_DOCX_DIR.exists(): return {}
    try:
        from docx import Document
        from difflib import get_close_matches
    except ImportError:
        print("  WARNING: python-docx not installed — roster parsing skipped.")
        return {}
    name_to_row = {m["name"]: m for m in mapping if m["name"]}
    names = list(name_to_row.keys())
    rosters = {}
    unmatched = []
    for f in sorted(SAMPLE_DOCX_DIR.glob("*.docx")):
        if f.name.startswith("~$"): continue   # Word lock files
        try:
            doc = Document(str(f))
        except Exception as e:
            print(f"  Skipping {f.name}: {e}"); continue
        if len(doc.tables) < 1: continue
        name = _docx_school_name(doc)
        if not name:
            unmatched.append((f.name, "no School Name cell found")); continue
        match = get_close_matches(name, names, n=1, cutoff=0.7)
        if not match:
            unmatched.append((f.name, f"no mapping match for '{name}'")); continue
        sid = name_to_row[match[0]]["school_id"]
        t1 = _docx_table_student_ids(doc.tables[0]) if len(doc.tables) >= 1 else set()
        t2 = _docx_table_student_ids(doc.tables[1]) if len(doc.tables) >= 2 else set()
        if sid in rosters:
            rosters[sid]["t1"] |= t1
            rosters[sid]["t2"] |= t2
            rosters[sid]["files"].append(f.name)
        else:
            rosters[sid] = {"t1": t1, "t2": t2, "files": [f.name]}
    if unmatched:
        print(f"  {len(unmatched)} roster file(s) unmatched:")
        for fn, reason in unmatched[:5]: print(f"    - {fn}: {reason}")
    return rosters

def _rows_by_school(rows):
    d = defaultdict(list)
    for r in rows:
        sid = _norm_id(r.get("school_id", ""))
        if sid.isdigit(): d[int(sid)].append(r)
    return d

def _girls_candidate_ids(row):
    """IDs from a girls endline row that should be checked against the roster."""
    return {i for i in (_norm_id(row.get("b_reg_id", "")),
                        _norm_id(row.get("new_reg_id", ""))) if i}

def _boys_candidate_ids(row):
    """IDs from a boys endline row. Boys CSV has only new_reg_id."""
    return {i for i in (_norm_id(row.get("new_reg_id", "")),) if i}

mapping = _load_school_mapping()
rosters = _build_rosters(mapping)

# Bucket the endline rows
girls_by_school = _rows_by_school(girls)
boys_by_school  = _rows_by_school(boys)

cov = {
    "girls_treatment": {"bl_tr": 0, "tr_only": 0, "new": 0,
                        "schools_loaded": 0, "schools_total": 0, "endline_rows_loaded": 0},
    "girls_control":   {"bl": 0, "new": 0,
                        "schools_loaded": 0, "schools_total": 0, "endline_rows_loaded": 0},
    "boys_treatment":  {"trained": 0, "new": 0,
                        "schools_loaded": 0, "schools_total": 0, "endline_rows_loaded": 0},
}

gt_per_school = []   # per-school detail for Girls Treatment (61 schools, target 40)
gc_per_school = []   # per-school detail for Girls Control   (29 schools, target 20)
bt_per_school = []   # per-school detail for Boys Treatment  (22 schools, target 20)

for m in mapping:
    cat = m["category"]
    if cat in (None, "boys_control"): continue
    cov[cat]["schools_total"] += 1
    sid = m["school_id"]
    roster = rosters.get(sid)

    if cat == "girls_treatment":
        rows = girls_by_school.get(sid, [])
        total = len(rows)
        if roster:
            cov[cat]["schools_loaded"] += 1
            cov[cat]["endline_rows_loaded"] += total
            t1, t2 = roster["t1"], roster["t2"]
            bl_tr = tr_only = new_ = 0
            for r in rows:
                cands = _girls_candidate_ids(r)
                if cands & t1:   bl_tr   += 1
                elif cands & t2: tr_only += 1
                else:            new_    += 1
            cov[cat]["bl_tr"]   += bl_tr
            cov[cat]["tr_only"] += tr_only
            cov[cat]["new"]     += new_
            gt_per_school.append({"sid": sid, "name": m["name"], "loaded": True,
                                  "total": total, "bl_tr": bl_tr, "tr_only": tr_only, "new": new_})
        else:
            gt_per_school.append({"sid": sid, "name": m["name"], "loaded": False,
                                  "total": total, "bl_tr": None, "tr_only": None, "new": None})

    elif cat == "girls_control":
        rows = girls_by_school.get(sid, [])
        total = len(rows)
        if roster:
            cov[cat]["schools_loaded"] += 1
            cov[cat]["endline_rows_loaded"] += total
            t1 = roster["t1"]
            bl = new_ = 0
            for r in rows:
                if _girls_candidate_ids(r) & t1: bl   += 1
                else:                            new_ += 1
            cov[cat]["bl"]  += bl
            cov[cat]["new"] += new_
            gc_per_school.append({"sid": sid, "name": m["name"], "loaded": True,
                                  "total": total, "bl": bl, "new": new_})
        else:
            gc_per_school.append({"sid": sid, "name": m["name"], "loaded": False,
                                  "total": total, "bl": None, "new": None})

    elif cat == "boys_treatment":
        rows = boys_by_school.get(sid, [])
        total = len(rows)
        if roster:
            cov[cat]["schools_loaded"] += 1
            cov[cat]["endline_rows_loaded"] += total
            t1 = roster["t1"]
            trained = new_ = 0
            for r in rows:
                if _boys_candidate_ids(r) & t1: trained += 1
                else:                           new_    += 1
            cov[cat]["trained"] += trained
            cov[cat]["new"]     += new_
            bt_per_school.append({"sid": sid, "name": m["name"], "loaded": True,
                                  "total": total, "trained": trained, "new": new_})
        else:
            bt_per_school.append({"sid": sid, "name": m["name"], "loaded": False,
                                  "total": total, "trained": None, "new": None})

# Convenience locals used in subs (prefixed with cov_)
gt = cov["girls_treatment"]; gc = cov["girls_control"]; bt = cov["boys_treatment"]
cov_gt_bl_tr   = gt["bl_tr"];   cov_gt_tr_only = gt["tr_only"]; cov_gt_new     = gt["new"]
cov_gt_loaded  = gt["schools_loaded"]; cov_gt_total = gt["schools_total"]; cov_gt_rows = gt["endline_rows_loaded"]
cov_gc_bl      = gc["bl"];      cov_gc_new     = gc["new"]
cov_gc_loaded  = gc["schools_loaded"]; cov_gc_total = gc["schools_total"]; cov_gc_rows = gc["endline_rows_loaded"]
cov_bt_trained = bt["trained"]; cov_bt_new     = bt["new"]
cov_bt_loaded  = bt["schools_loaded"]; cov_bt_total = bt["schools_total"]; cov_bt_rows = bt["endline_rows_loaded"]

# Expected-survey denominators (project design):
#   Girls treatment: 40 endline / school (20 baseline+training + 20 training-only)
#   Girls control  : 20 endline / school
#   Boys treatment : 20 endline / school (assumed — revise when first roster lands)
cov_gt_target = cov_gt_loaded * 40
cov_gc_target = cov_gc_loaded * 20
cov_bt_target = cov_bt_loaded * 20
cov_gt_pct    = pct(cov_gt_rows, cov_gt_target)
cov_gc_pct    = pct(cov_gc_rows, cov_gc_target)
cov_bt_pct    = pct(cov_bt_rows, cov_bt_target)

# Boys treatment state — show "awaiting rosters" if none loaded yet
cov_bt_has_rosters = cov_bt_loaded > 0

# Per-school table rows — one HTML block per arm
import html as _html

COMPLETION_THRESHOLD = 15   # School is "completed" if endline surveys >= 15

def _build_school_rows(per_school, endline_target, bucket_keys, pct_key, pct_denom):
    """Build HTML <tr>…</tr> rows for completed schools only.
    A school qualifies if its roster is loaded AND endline total >= COMPLETION_THRESHOLD.
    Returns an empty-state row if no school qualifies."""
    # Column count = sid + name + endline + buckets + %
    ncols = 3 + len(bucket_keys) + 1
    out = []
    for s in sorted(per_school, key=lambda x: x["sid"]):
        if not s["loaded"] or s["total"] < COMPLETION_THRESHOLD:
            continue
        name_esc = _html.escape(s["name"])
        row_pct = pct(s[pct_key], pct_denom)
        bucket_cells = "".join(
            f"<td style='text-align:right;font-weight:700;'>{s[k]}</td>"
            for k in bucket_keys
        )
        out.append(
            f"<tr><td style='text-align:right;'>{s['sid']}</td>"
            f"<td>{name_esc}</td>"
            f"<td style='text-align:right;'>{s['total']}"
            f"<span style='color:var(--rs-navy-soft);'> / {endline_target}</span></td>"
            f"{bucket_cells}"
            f"<td style='text-align:right;'>{row_pct}%</td></tr>"
        )
    if not out:
        out.append(
            f"<tr><td colspan='{ncols}' style='text-align:center;padding:24px;"
            f"color:var(--rs-navy-soft);font-style:italic;'>"
            f"No schools completed yet — endline collection in progress.</td></tr>"
        )
    return "\n          ".join(out)

# Girls Treatment: endline target per school = 40 (20 BL+TR + 20 TR-only).
# Percentage tracks the panel-recapture rate: BL+TR re-found out of the 20 baseline girls.
cov_gt_school_rows = _build_school_rows(
    gt_per_school, endline_target=40, bucket_keys=["bl_tr", "tr_only", "new"],
    pct_key="bl_tr", pct_denom=20,
)
# Girls Control: target 20 baseline girls; % = BL re-found out of 20.
cov_gc_school_rows = _build_school_rows(
    gc_per_school, endline_target=20, bucket_keys=["bl", "new"],
    pct_key="bl", pct_denom=20,
)
# Boys Treatment: target 20 trained boys; % = Trained re-found out of 20.
cov_bt_school_rows = _build_school_rows(
    bt_per_school, endline_target=20, bucket_keys=["trained", "new"],
    pct_key="trained", pct_denom=20,
)

print(f"  Girls Treatment : {cov_gt_loaded}/{cov_gt_total} schools loaded, "
      f"{cov_gt_rows} endline rows -> BL+TR={cov_gt_bl_tr}, TR-only={cov_gt_tr_only}, New={cov_gt_new}")
print(f"  Girls Control   : {cov_gc_loaded}/{cov_gc_total} schools loaded, "
      f"{cov_gc_rows} endline rows -> BL={cov_gc_bl}, New={cov_gc_new}")
print(f"  Boys Treatment  : {cov_bt_loaded}/{cov_bt_total} schools loaded, "
      f"{cov_bt_rows} endline rows -> Trained={cov_bt_trained}, New={cov_bt_new}")

# ══════════════════════════════════════════════════════════════════════════
# BUILD JAVASCRIPT CHART BLOCK
# ══════════════════════════════════════════════════════════════════════════
print("Generating chart data block...")

# Stacked safety detail — 4 response categories × 6 bars
safe_d1 = [boys_commute_dist[0], girls_commute_dist[0], boys_school_dist[0], girls_school_dist[0], boys_home_dist[0], girls_home_dist[0]]
safe_d2 = [boys_commute_dist[1], girls_commute_dist[1], boys_school_dist[1], girls_school_dist[1], boys_home_dist[1], girls_home_dist[1]]
safe_d3 = [boys_commute_dist[2], girls_commute_dist[2], boys_school_dist[2], girls_school_dist[2], boys_home_dist[2], girls_home_dist[2]]
safe_d4 = [boys_commute_dist[3], girls_commute_dist[3], boys_school_dist[3], girls_school_dist[3], boys_home_dist[3], girls_home_dist[3]]

# Stacked gender norms detail — labels: G2B,G2G,G3B,G3G,G5B,G5G,G6B,G6G
def gn_stacked(level):  # level = 0..4 (strongly disagree .. strongly agree)
    return [boys_gn_dists[i][level] for i in range(4) for _ in (0,1)][::-1]

# Interleave boys/girls for each of 4 gn items
gn_boys_girls = []
for i in range(4):
    gn_boys_girls.extend([boys_gn_dists[i], girls_gn_dists[i]])

gn_sd = [gn_boys_girls[j][0] for j in range(8)]
gn_d  = [gn_boys_girls[j][1] for j in range(8)]
gn_n  = [gn_boys_girls[j][2] for j in range(8)]
gn_a  = [gn_boys_girls[j][3] for j in range(8)]
gn_sa = [gn_boys_girls[j][4] for j in range(8)]

# Stacked masculinity detail — 5 items × 5 levels
masc_sd = [masc_dists[i][0] for i in range(5)]
masc_d  = [masc_dists[i][1] for i in range(5)]
masc_n  = [masc_dists[i][2] for i in range(5)]
masc_a  = [masc_dists[i][3] for i in range(5)]
masc_sa = [masc_dists[i][4] for i in range(5)]

# PQL distributions per response level (5 levels × 5 items)
def pql_level(dists, level):
    return [dists[i][level] for i in range(5)]

bpql = [pql_level(boys_pql_dists,  l) for l in range(5)]
gpql = [pql_level(girls_pql_dists, l) for l in range(5)]

js = f"""// ══════════════════════════════════════════
// AUTO-GENERATED — {last_updated}
// Boys: {n_boys} records  |  Girls: {n_girls} records  |  Total: {n_total}
// ══════════════════════════════════════════

// ══════════════════════════════════════════
// OVERVIEW CHARTS
// ══════════════════════════════════════════

new Chart(document.getElementById('chartGenderSplit'), {{
  type: 'doughnut',
  data: {{
    labels: ['Boys ({n_boys})', 'Girls ({n_girls})'],
    datasets: [{{ data: [{n_boys}, {n_girls}], backgroundColor: [NAVY, RED], borderWidth: 3, borderColor: '#fff' }}]
  }},
  options: {{
    ...defaults,
    plugins: {{
      ...defaults.plugins,
      datalabels: {{ display: true, color: '#fff', font: {{ family: 'Inter', size: 13, weight: 700 }}, formatter: v => v + ' students' }},
      legend: {{ position: 'bottom', labels: {{ font: {{ family: 'Inter', size: 13 }}, padding: 16 }} }}
    }}
  }}
}});

new Chart(document.getElementById('chartRadar'), {{
  type: 'radar',
  data: {{
    labels: ['Academic\\nPerformance', 'Safety\\nPerception', 'Wellbeing\\n(PQL)', 'Self-\\nEfficacy', 'Gender\\nEquity Norms'],
    datasets: [
      {{ label: 'Boys',  data: {js_list(boys_radar)},  borderColor: NAVY, backgroundColor: NAVY_L, borderWidth: 2, pointBackgroundColor: NAVY, pointRadius: 5 }},
      {{ label: 'Girls', data: {js_list(girls_radar)}, borderColor: RED,  backgroundColor: RED_L,  borderWidth: 2, pointBackgroundColor: RED,  pointRadius: 5 }}
    ]
  }},
  options: {{
    ...defaults,
    plugins: {{ ...defaults.plugins, datalabels: {{ display: false }}, legend: {{ position: 'bottom' }} }},
    scales: {{ r: {{ min: 0, max: 100, ticks: {{ font: {{ family: 'Inter', size: 10 }}, color: '#475569', stepSize: 25 }}, grid: {{ color: '#e2e8f0' }}, pointLabels: {{ font: {{ family: 'Inter', size: 11 }}, color: '#334155' }} }} }}
  }}
}});

new Chart(document.getElementById('chartAcadOverview'), {{
  type: 'bar',
  data: {{
    labels: ['Boys (n={n_boys})', 'Girls (n={n_girls})'],
    datasets: [{{ label: 'Avg Grade 7 %', data: [{boys_grade_mean}, {girls_grade_mean}], backgroundColor: [NAVY, RED], borderRadius: 8, borderSkipped: false }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 13, weight: 700 }}, formatter: v => v + '%' }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 80, title: {{ display: true, text: 'Average %' }} }} }} }}
}});

new Chart(document.getElementById('chartSafetyOverview'), {{
  type: 'bar',
  data: {{
    labels: ['Commute\\nto School', 'School\\nNeighbourhood', 'Home\\nArea'],
    datasets: [
      {{ label: 'Boys',  data: [{boys_commute_mean},  {boys_school_mean},  {boys_home_mean}],  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: [{girls_commute_mean}, {girls_school_mean}, {girls_home_mean}], backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 11, weight: 600 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 4.5, title: {{ display: true, text: 'Mean Score (1–4)' }} }} }} }}
}});

new Chart(document.getElementById('chartAspOverview'), {{
  type: 'doughnut',
  data: {{
    labels: ['Matric', 'Intermediate', 'Bachelors', 'Masters', 'Professional Degree'],
    datasets: [{{ label: 'Girls Aspirations', data: {js_list(girls_asp)}, backgroundColor: [RED, AMBER, TEAL, NAVY, PURPLE], borderWidth: 3, borderColor: '#fff' }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, color: '#fff', font: {{ family: 'Inter', size: 11, weight: 700 }}, formatter: (v, ctx) => v > 0 ? v : '' }}, legend: {{ position: 'bottom', labels: {{ font: {{ family: 'Inter', size: 11 }}, padding: 10 }} }}, title: {{ display: true, text: 'Girls (n={n_girls})', font: {{ family: 'Inter', size: 12 }}, color: '#6B7494', padding: {{ bottom: 6 }} }} }} }}
}});

// ══════════════════════════════════════════
// STUDENT PROFILE CHARTS
// ══════════════════════════════════════════

new Chart(document.getElementById('chartFatherEdu'), {{
  type: 'bar',
  data: {{
    labels: ['None/Primary', 'Middle', 'Matric', 'Intermediate', 'Bachelors', 'Masters+'],
    datasets: [
      {{ label: 'Boys', data: {js_list(boys_father_edu)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: {js_list(girls_father_edu)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

new Chart(document.getElementById('chartMotherEdu'), {{
  type: 'bar',
  data: {{
    labels: ['None/Primary', 'Middle', 'Matric', 'Intermediate', 'Masters+'],
    datasets: [
      {{ label: 'Boys', data: {js_list(boys_mother_edu)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: {js_list(girls_mother_edu)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

new Chart(document.getElementById('chartFatherOcc'), {{
  type: 'bar',
  data: {{
    labels: ['Government\\nEmployee', 'Private\\nEmployee', 'Self-Employed\\n/ Business', 'Farming', 'Skilled\\nLabour', 'Other'],
    datasets: [
      {{ label: 'Boys', data: {js_list(boys_father_occ)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: {js_list(girls_father_occ)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

new Chart(document.getElementById('chartMotherOcc'), {{
  type: 'bar',
  data: {{
    labels: ['Government\\nEmployee', 'Private\\nEmployee', 'Self-Employed', 'Homemaker', 'Other'],
    datasets: [
      {{ label: 'Boys', data: {js_list(boys_mother_occ)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: {js_list(girls_mother_occ)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

new Chart(document.getElementById('chartAssets'), {{
  type: 'bar',
  data: {{
    labels: ['Refrigerator', 'Television', 'Motorcycle/\\nBicycle', 'Car/Van', 'Computer/\\nLaptop', 'Own House', 'Smartphone', 'Internet\\n/WiFi', 'Agricultural\\nLand', 'Livestock', 'Bank\\nAccount', 'Business\\n/Shop', 'Jewellery\\n/Invest.'],
    datasets: [
      {{ label: 'Boys %',  data: {js_list(boys_assets_pct)},  backgroundColor: NAVY, borderRadius: 4, borderSkipped: false }},
      {{ label: 'Girls %', data: {js_list(girls_assets_pct)}, backgroundColor: RED,  borderRadius: 4, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, indexAxis: 'y', plugins: {{ ...defaults.plugins }}, scales: {{ x: {{ ...defaults.scales.x, min: 0, max: 110, title: {{ display: true, text: '% Owning Asset' }} }}, y: {{ ticks: {{ font: {{ family: 'Inter', size: 10 }}, color: '#475569' }}, grid: {{ display: false }} }} }} }}
}});

new Chart(document.getElementById('chartLanguage'), {{
  type: 'doughnut',
  data: {{
    labels: ['Urdu', 'Punjabi', 'Other'],
    datasets: [{{ label: 'Boys', data: {js_list(boys_lang)}, backgroundColor: [NAVY, TEAL, '#94a3b8'], borderWidth: 3, borderColor: '#fff' }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, color: '#fff', font: {{ family: 'Inter', size: 12, weight: 700 }}, formatter: (v, ctx) => v > 0 ? ctx.chart.data.labels[ctx.dataIndex] + '\\n' + v : '' }}, legend: {{ position: 'bottom' }}, title: {{ display: true, text: 'Boys (n={n_boys}) — Girls similar', font: {{ family: 'Inter', size: 12 }}, color: '#6B7494' }} }} }}
}});

// ══════════════════════════════════════════
// ACADEMIC CHARTS
// ══════════════════════════════════════════

new Chart(document.getElementById('chartGradeDistribution'), {{
  type: 'bar',
  data: {{
    labels: ['30–39%', '40–49%', '50–59%', '60–69%', '70–79%', '80%+'],
    datasets: [
      {{ label: 'Boys (n={n_boys})',  data: {js_list(boys_grade_dist)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls (n={n_girls})', data: {js_list(girls_grade_dist)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 11, weight: 600 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

new Chart(document.getElementById('chartAspirations'), {{
  type: 'bar',
  data: {{
    labels: ['Matric\\n(Gr.10)', 'Intermediate\\n(Gr.12)', 'Bachelors', 'Masters', 'Professional\\nDegree'],
    datasets: [
      {{ label: 'Boys (n={n_boys})',  data: {js_list(boys_asp)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls (n={n_girls})', data: {js_list(girls_asp)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 11, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

new Chart(document.getElementById('chartAbsence'), {{
  type: 'bar',
  data: {{
    labels: ['Boys', 'Girls'],
    datasets: [{{ label: 'Avg Absence Days', data: [{boys_abs_mean}, {girls_abs_mean}], backgroundColor: [NAVY, RED], borderRadius: 8, borderSkipped: false }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 14, weight: 700 }}, formatter: v => v + ' days' }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: Math.max({boys_abs_mean}, {girls_abs_mean}) + 2, title: {{ display: true, text: 'Average Days Absent' }} }} }} }}
}});

new Chart(document.getElementById('chartPerfBands'), {{
  type: 'bar',
  data: {{
    labels: ['Below 45%', '45–59%', '60–74%', '75%+'],
    datasets: [
      {{ label: 'Boys',  data: {js_list(boys_perf_bands)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: {js_list(girls_perf_bands)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 11, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

// ══════════════════════════════════════════
// SAFETY CHARTS
// ══════════════════════════════════════════

new Chart(document.getElementById('chartSafetyDetail'), {{
  type: 'bar',
  data: {{
    labels: ['Commute — Boys', 'Commute — Girls', 'School Area — Boys', 'School Area — Girls', 'Home Area — Boys', 'Home Area — Girls'],
    datasets: [
      {{ label: '1 — Not Safe',       data: {js_list(safe_d1)}, backgroundColor: '#E8192C', borderSkipped: false }},
      {{ label: '2 — Not Very Safe',  data: {js_list(safe_d2)}, backgroundColor: '#F59E0B', borderSkipped: false }},
      {{ label: '3 — Somewhat Safe',  data: {js_list(safe_d3)}, backgroundColor: '#0891B2', borderSkipped: false }},
      {{ label: '4 — Very Safe',      data: {js_list(safe_d4)}, backgroundColor: '#2D3450', borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, indexAxis: 'y', plugins: {{ ...defaults.plugins }}, scales: {{ x: {{ stacked: true, ...defaults.scales.x, title: {{ display: true, text: 'Number of Students' }} }}, y: {{ stacked: true, ticks: {{ font: {{ family: 'Inter', size: 11 }}, color: '#475569' }}, grid: {{ display: false }} }} }} }}
}});

new Chart(document.getElementById('chartTransport'), {{
  type: 'bar',
  data: {{
    labels: ['Walking', 'Vehicle\\n(Van/Car/Motorcycle)', 'Other'],
    datasets: [
      {{ label: 'Boys (n={n_boys})',  data: {js_list(boys_transport)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls (n={n_girls})', data: {js_list(girls_transport + [0])}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 12, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

new Chart(document.getElementById('chartTravelTime'), {{
  type: 'bar',
  data: {{
    labels: ['Boys', 'Girls'],
    datasets: [{{ label: 'Avg Travel Time (minutes)', data: [{boys_travel_mean}, {girls_travel_mean}], backgroundColor: [NAVY, RED], borderRadius: 8, borderSkipped: false }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 14, weight: 700 }}, formatter: v => v + ' min' }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: Math.max({boys_travel_mean}, {girls_travel_mean}) + 5, title: {{ display: true, text: 'Minutes' }} }} }} }}
}});

new Chart(document.getElementById('chartBehaviorMods'), {{
  type: 'bar',
  data: {{
    labels: ['Changed\\nRoute', 'Changed\\nTiming', 'Changed\\nClothing', 'Identifies\\nEscape Route'],
    datasets: [
      {{ label: 'Always',    data: [{girls_route[0]}, {girls_timing[0]}, {girls_clothing[0]}, {girls_escape[0]}], backgroundColor: '#E8192C', borderRadius: [4,4,0,0], borderSkipped: false }},
      {{ label: 'Sometimes', data: [{girls_route[1]}, {girls_timing[1]}, {girls_clothing[1]}, {girls_escape[1]}], backgroundColor: '#F59E0B', borderRadius: [4,4,0,0], borderSkipped: false }},
      {{ label: 'Rarely',    data: [{girls_route[2]}, {girls_timing[2]}, {girls_clothing[2]}, {girls_escape[2]}], backgroundColor: '#0891B2', borderRadius: [4,4,0,0], borderSkipped: false }},
      {{ label: 'Never',     data: [{girls_route[3]}, {girls_timing[3]}, {girls_clothing[3]}, {girls_escape[3]}], backgroundColor: '#2D3450', borderRadius: [4,4,0,0], borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins }}, scales: {{ x: {{ stacked: true, ...defaults.scales.x }}, y: {{ stacked: true, ...defaults.scales.y, title: {{ display: true, text: 'Girls (n={n_girls})' }} }} }} }}
}});

// ══════════════════════════════════════════
// GENDER NORMS CHARTS
// ══════════════════════════════════════════

new Chart(document.getElementById('chartGenderNorms'), {{
  type: 'bar',
  data: {{
    labels: ['Girls: Focus\\non Household', 'Education\\nEqually Valued', 'Mothers\\nWorking OK', 'Men Must\\nProtect Women'],
    datasets: [
      {{ label: 'Boys',  data: {js_list(boys_gn_means)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: {js_list(girls_gn_means)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 11, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 6, title: {{ display: true, text: 'Mean Score (1–5)' }} }} }} }}
}});

new Chart(document.getElementById('chartGenderNormsDetail'), {{
  type: 'bar',
  data: {{
    labels: ['G2 Boys', 'G2 Girls', 'G3 Boys', 'G3 Girls', 'G5 Boys', 'G5 Girls', 'G6 Boys', 'G6 Girls'],
    datasets: [
      {{ label: 'Strongly Disagree (1)', data: {js_list(gn_sd)}, backgroundColor: '#0891B2' }},
      {{ label: 'Disagree (2)',          data: {js_list(gn_d)},  backgroundColor: '#67E8F9' }},
      {{ label: 'Neutral (3)',           data: {js_list(gn_n)},  backgroundColor: '#E2E8F0' }},
      {{ label: 'Agree (4)',             data: {js_list(gn_a)},  backgroundColor: '#F59E0B' }},
      {{ label: 'Strongly Agree (5)',    data: {js_list(gn_sa)}, backgroundColor: '#E8192C' }}
    ]
  }},
  options: {{ ...defaults, indexAxis: 'y', plugins: {{ ...defaults.plugins }}, scales: {{ x: {{ stacked: true, ...defaults.scales.x, title: {{ display: true, text: 'Number of Students' }} }}, y: {{ stacked: true, ticks: {{ font: {{ family: 'Inter', size: 11 }} }}, grid: {{ display: false }} }} }} }}
}});

new Chart(document.getElementById('chartHarassWitness'), {{
  type: 'bar',
  data: {{
    labels: ['Physical Harassment\\n(Witnessed)', 'Verbal Harassment\\n(Witnessed)'],
    datasets: [
      {{ label: 'Boys — Yes (witnessed)',  data: [{boys_phys_witness},  {boys_verbal_witness}],  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls — Yes (witnessed)', data: [{girls_phys_witness}, {girls_verbal_witness}], backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 13, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, max: {max(n_boys,n_girls)+5}, title: {{ display: true, text: 'Number of Students Saying "Yes"' }} }} }} }}
}});

new Chart(document.getElementById('chartBlame'), {{
  type: 'bar',
  data: {{
    labels: ['1 — Victim\\'s\\nFault', '2 — More Victim\\nThan Perp.', '3 — Both\\nEqually', '4 — More Perp.\\nThan Victim', '5 — Perp.\\nFully at Fault'],
    datasets: [
      {{ label: 'Boys (n={n_boys})',  data: {js_list(boys_blame)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls (n={n_girls})', data: {js_list(girls_blame)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 11, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, title: {{ display: true, text: 'Number of Students' }} }} }} }}
}});

// ══════════════════════════════════════════
// WELLBEING CHARTS
// ══════════════════════════════════════════

new Chart(document.getElementById('chartPQL'), {{
  type: 'bar',
  data: {{
    labels: ['PQL2\\nMood/Feelings', 'PQL3\\nSchool/Learning', 'PQL5\\nFriendships', 'PQL6\\nFamily', 'PQL15\\nOverall Life'],
    datasets: [
      {{ label: 'Boys',  data: {js_list(boys_pql_means)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: {js_list(girls_pql_means)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 11, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 5.5, title: {{ display: true, text: 'Mean Score (1–5)' }} }} }} }}
}});

new Chart(document.getElementById('chartSelfEfficacy'), {{
  type: 'bar',
  data: {{
    labels: ['SE1: Can Handle\\nChallenges', 'SE2: Can Find\\nSolutions', 'SE3: Stay Calm\\nUnder Pressure'],
    datasets: [
      {{ label: 'Boys',  data: {js_list(boys_se_means)},  backgroundColor: NAVY, borderRadius: 6, borderSkipped: false }},
      {{ label: 'Girls', data: {js_list(girls_se_means)}, backgroundColor: RED,  borderRadius: 6, borderSkipped: false }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 12, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 3.5, title: {{ display: true, text: 'Mean Score (1–3)' }} }} }} }}
}});

new Chart(document.getElementById('chartPQLGirls'), {{
  type: 'bar',
  data: {{
    labels: ['PQL2\\nMood', 'PQL3\\nLearning', 'PQL5\\nFriendships', 'PQL6\\nFamily', 'PQL15\\nOverall'],
    datasets: [
      {{ label: 'Never (1)',        data: {js_list(gpql[0])}, backgroundColor: '#E8192C' }},
      {{ label: 'Almost Never (2)', data: {js_list(gpql[1])}, backgroundColor: '#F59E0B' }},
      {{ label: 'Sometimes (3)',    data: {js_list(gpql[2])}, backgroundColor: '#E2E8F0' }},
      {{ label: 'Often (4)',        data: {js_list(gpql[3])}, backgroundColor: '#0891B2' }},
      {{ label: 'Almost Always (5)',data: {js_list(gpql[4])}, backgroundColor: '#2D3450' }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins }}, scales: {{ x: {{ stacked: true, ...defaults.scales.x }}, y: {{ stacked: true, ...defaults.scales.y, title: {{ display: true, text: 'Girls (n={n_girls})' }} }} }} }}
}});

new Chart(document.getElementById('chartPQLBoys'), {{
  type: 'bar',
  data: {{
    labels: ['PQL2\\nMood', 'PQL3\\nLearning', 'PQL5\\nFriendships', 'PQL6\\nFamily', 'PQL15\\nOverall'],
    datasets: [
      {{ label: 'Never (1)',        data: {js_list(bpql[0])}, backgroundColor: '#E8192C' }},
      {{ label: 'Almost Never (2)', data: {js_list(bpql[1])}, backgroundColor: '#F59E0B' }},
      {{ label: 'Sometimes (3)',    data: {js_list(bpql[2])}, backgroundColor: '#E2E8F0' }},
      {{ label: 'Often (4)',        data: {js_list(bpql[3])}, backgroundColor: '#0891B2' }},
      {{ label: 'Almost Always (5)',data: {js_list(bpql[4])}, backgroundColor: '#2D3450' }}
    ]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins }}, scales: {{ x: {{ stacked: true, ...defaults.scales.x }}, y: {{ stacked: true, ...defaults.scales.y, title: {{ display: true, text: 'Boys (n={n_boys})' }} }} }} }}
}});

// ══════════════════════════════════════════
// BOYS' MODULE CHARTS
// ══════════════════════════════════════════

new Chart(document.getElementById('chartMasculinity'), {{
  type: 'bar',
  data: {{
    labels: ['M1: Men Do\\nRepairs', 'M2: Men Don\\'t\\nAsk Help', 'M3: Men Don\\'t\\nShow Fear', 'M4: Must Win\\nGames', 'M5: Must Tease\\nto Fit In'],
    datasets: [{{ label: 'Mean Score (Boys, n={n_boys})', data: {js_list(masc_means)}, backgroundColor: [NAVY, NAVY, AMBER, AMBER, TEAL], borderRadius: 8, borderSkipped: false }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 13, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 5.5, title: {{ display: true, text: 'Mean Agreement (1–5)' }} }} }} }}
}});

new Chart(document.getElementById('chartMascDetail'), {{
  type: 'bar',
  data: {{
    labels: ['M1: Men Do Repairs', 'M2: Don\\'t Ask Help', 'M3: Hide Fear', 'M4: Must Win', 'M5: Tease to Fit In'],
    datasets: [
      {{ label: 'Strongly Disagree (1)', data: {js_list(masc_sd)}, backgroundColor: '#0891B2' }},
      {{ label: 'Disagree (2)',          data: {js_list(masc_d)},  backgroundColor: '#67E8F9' }},
      {{ label: 'Neutral (3)',           data: {js_list(masc_n)},  backgroundColor: '#E2E8F0' }},
      {{ label: 'Agree (4)',             data: {js_list(masc_a)},  backgroundColor: '#F59E0B' }},
      {{ label: 'Strongly Agree (5)',    data: {js_list(masc_sa)}, backgroundColor: '#E8192C' }}
    ]
  }},
  options: {{ ...defaults, indexAxis: 'y', plugins: {{ ...defaults.plugins }}, scales: {{ x: {{ stacked: true, ...defaults.scales.x, title: {{ display: true, text: 'Number of Boys' }} }}, y: {{ stacked: true, ticks: {{ font: {{ family: 'Inter', size: 11 }} }}, grid: {{ display: false }} }} }} }}
}});

new Chart(document.getElementById('chartImpulse'), {{
  type: 'bar',
  data: {{
    labels: ['I1', 'I2', 'I3', 'I4', 'I5', 'I6', 'I7', 'I8'],
    datasets: [{{ label: 'Mean Score', data: {js_list(impulse_means)}, backgroundColor: [NAVY, NAVY, NAVY, AMBER, RED, RED, AMBER, TEAL], borderRadius: 6, borderSkipped: false }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 11, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 5.5, title: {{ display: true, text: 'Mean Score (1=Never, 5=Always)' }} }} }} }}
}});

new Chart(document.getElementById('chartCBTPeer'), {{
  type: 'bar',
  data: {{
    labels: ['CBT1', 'CBT2', 'CBT3', 'Peer 1', 'Peer 2', 'Peer 3'],
    datasets: [{{ label: 'Mean Score', data: {js_list(cbt_means + peer_means)}, backgroundColor: [PURPLE, PURPLE, PURPLE, AMBER, AMBER, AMBER], borderRadius: 6, borderSkipped: false }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 12, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 4.5, title: {{ display: true, text: 'Mean Score' }} }} }} }}
}});

// ══════════════════════════════════════════
// GIRLS' MODULE CHARTS
// ══════════════════════════════════════════

new Chart(document.getElementById('chartTimeUse'), {{
  type: 'bar',
  data: {{
    labels: ['Fetching\\nWater', 'Cooking', 'Cleaning', 'Caring\\n(Children/Elderly)', 'Business\\n/ Work'],
    datasets: [{{ label: 'Avg Daily Hours', data: {js_list(girls_time_use)}, backgroundColor: [TEAL, AMBER, RED, NAVY, PURPLE], borderRadius: 8, borderSkipped: false }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 13, weight: 700 }}, formatter: v => v + 'h' }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, min: 0, max: 5, title: {{ display: true, text: 'Average Hours per Day (Girls, n={n_girls})' }} }} }} }}
}});

new Chart(document.getElementById('chartGirlsBehavior'), {{
  type: 'doughnut',
  data: {{
    labels: ['Always Change Clothing', 'Sometimes Change Clothing', 'Rarely', 'Never Change'],
    datasets: [{{ data: {js_list(girls_clothing)}, backgroundColor: [RED, AMBER, TEAL, NAVY], borderWidth: 3, borderColor: '#fff' }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, color: '#fff', font: {{ family: 'Inter', size: 12, weight: 700 }}, formatter: (v, ctx) => v > 0 ? v : '' }}, legend: {{ position: 'bottom', labels: {{ font: {{ family: 'Inter', size: 11 }}, padding: 10 }} }}, title: {{ display: true, text: 'Clothing Modification (n={n_girls} girls)', font: {{ family: 'Inter', size: 12 }}, color: '#6B7494' }} }} }}
}});

new Chart(document.getElementById('chartHarassShare'), {{
  type: 'bar',
  data: {{
    labels: ['Would Tell\\nNo One', 'Would Tell\\nFriend/Sibling', 'Would Tell\\nParent/Family'],
    datasets: [{{ label: 'Girls Who Would Tell (n={n_girls})', data: [{girls_share_nobody}, {girls_share_friend}, {girls_share_parent}], backgroundColor: [RED, TEAL, NAVY], borderRadius: 8, borderSkipped: false }}]
  }},
  options: {{ ...defaults, plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 14, weight: 700 }} }} }}, scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, max: {n_girls+3}, title: {{ display: true, text: 'Number of Girls' }} }} }} }}
}});

new Chart(document.getElementById('chartHarassGender'), {{
  type: 'bar',
  data: {{
    labels: ['Boys', 'Girls'],
    datasets: [{{ label: 'Witnessed Verbal Harassment', data: [{boys_verbal_witness}, {girls_verbal_witness}], backgroundColor: [NAVY, RED], borderRadius: 8, borderSkipped: false }}]
  }},
  options: {{
    ...defaults,
    plugins: {{ ...defaults.plugins, datalabels: {{ display: true, anchor: 'end', align: 'end', color: '#334155', font: {{ family: 'Inter', size: 14, weight: 700 }}, formatter: (v, ctx) => {{ const t = [{n_boys},{n_girls}]; return v + ' / ' + t[ctx.dataIndex] + ' (' + Math.round(v/t[ctx.dataIndex]*100) + '%)'; }} }} }},
    scales: {{ ...defaults.scales, y: {{ ...defaults.scales.y, max: {max(n_boys,n_girls)+5}, title: {{ display: true, text: 'Number of Students' }} }} }}
  }}
}});
"""

# ══════════════════════════════════════════════════════════════════════════
# INJECT INTO TEMPLATE & WRITE index.html
# ══════════════════════════════════════════════════════════════════════════
print("Reading template...")
template_text = TEMPLATE.read_text(encoding="utf-8")

if "// %%CHART_BLOCK%%" not in template_text:
    raise RuntimeError("Placeholder '// %%CHART_BLOCK%%' not found in template.html")

output_text = template_text.replace("// %%CHART_BLOCK%%", js)

# Substitute KPI placeholders in the Overview section
subs = {
    "{n_total}":   n_total,   "{n_boys}":    n_boys,    "{n_girls}":   n_girls,   "{n_schools}": n_schools,
    "{boys_age_mean}":  boys_age_mean,  "{girls_age_mean}": girls_age_mean,
    "{boys_age_min}":   boys_age_min,   "{boys_age_max}":   boys_age_max,
    "{girls_age_min}":  girls_age_min,  "{girls_age_max}":  girls_age_max,
    "{boys_sib_mean}":  boys_sib_mean,  "{girls_sib_mean}": girls_sib_mean,
    "{boys_home_own_pct}":  boys_home_own_pct,  "{girls_home_own_pct}": girls_home_own_pct,
    "{boys_grade_mean}":    boys_grade_mean,    "{girls_grade_mean}":   girls_grade_mean,
    "{boys_grade_min}":     boys_grade_min,     "{boys_grade_max}":     boys_grade_max,
    "{girls_grade_min}":    girls_grade_min,    "{girls_grade_max}":    girls_grade_max,
    "{grade_gap}":          abs(grade_gap),
    "{grade_leader}":       "Boys" if boys_grade_mean >= girls_grade_mean else "Girls",
    "{grade_trailer}":      "girls" if boys_grade_mean >= girls_grade_mean else "boys",
    "{boys_abs_mean}":      boys_abs_mean,      "{girls_abs_mean}":     girls_abs_mean,
    "{boys_asp_pro_pct}":   boys_asp_pro_pct,   "{girls_asp_pro_pct}":  girls_asp_pro_pct,
    "{boys_commute_mean}":  boys_commute_mean,  "{girls_commute_mean}": girls_commute_mean,
    "{boys_school_mean}":   boys_school_mean,   "{girls_school_mean}":  girls_school_mean,
    "{boys_home_mean}":     boys_home_mean,     "{girls_home_mean}":    girls_home_mean,
    "{girls_clothing_mod_pct}": girls_clothing_mod_pct,
    "{girls_route_mod_pct}":    girls_route_mod_pct,
    "{girls_timing_mod_pct}":   girls_timing_mod_pct,
    "{girls_clothing_mod_n}":   girls_clothing_mod_n,
    "{girls_escape_pct}":       girls_escape_pct,
    "{boys_verbal_pct}":    boys_verbal_pct,    "{girls_verbal_pct}":   girls_verbal_pct,
    "{boys_blame_perp_pct}": boys_blame_perp_pct, "{girls_blame_perp_pct}": girls_blame_perp_pct,
    "{boys_pql_overall}":   boys_pql_overall,   "{girls_pql_overall}":  girls_pql_overall,
    "{boys_pql_family}":    boys_pql_means[1],  "{girls_pql_family}":   girls_pql_means[1],
    "{boys_pql_friend}":    boys_pql_means[2],  "{girls_pql_friend}":   girls_pql_means[2],
    "{boys_se_1}": boys_se_means[0], "{boys_se_2}": boys_se_means[1], "{boys_se_3}": boys_se_means[2],
    "{girls_se_1}": girls_se_means[0], "{girls_se_2}": girls_se_means[1], "{girls_se_3}": girls_se_means[2],
    "{boys_gn_household}": boys_gn_means[0], "{girls_gn_household}": girls_gn_means[0],
    "{boys_gn_protect}":   boys_gn_means[3], "{girls_gn_protect}":   girls_gn_means[3],
    "{masc_1}": masc_means[0], "{masc_2}": masc_means[1], "{masc_3}": masc_means[2], "{masc_4}": masc_means[3], "{masc_5}": masc_means[4],
    "{masc_1_pct}": round(masc_means[0]/5*100, 1),
    "{masc_2_pct}": round(masc_means[1]/5*100, 1),
    "{masc_3_pct}": round(masc_means[2]/5*100, 1),
    "{masc_4_pct}": round(masc_means[3]/5*100, 1),
    "{masc_5_pct}": round(masc_means[4]/5*100, 1),
    "{impulse_8}": impulse_means[7],
    "{peer_1}":    peer_means[0],
    "{time_cooking}": girls_time_use[1], "{time_cleaning}": girls_time_use[2], "{time_caring}": girls_time_use[3],
    "{girls_share_parent_pct}": girls_share_parent_pct,
    "{girls_share_friend_pct}": girls_share_friend_pct,
    "{girls_share_nobody_pct}": girls_share_nobody_pct,
    "{girls_report_pct}":       girls_report_pct,
    "{girls_report_n}":         girls_report,
    "{boys_father_inter_pct}":  boys_father_inter_pct,  "{girls_father_inter_pct}": girls_father_inter_pct,
    "{boys_father_lowed_pct}":  boys_father_lowed_pct,  "{girls_father_lowed_pct}": girls_father_lowed_pct,
    "{boys_mother_home_pct}":   boys_mother_home_pct,   "{girls_mother_home_pct}":  girls_mother_home_pct,
    # Sample Coverage tab
    "{cov_gt_bl_tr}":   cov_gt_bl_tr,   "{cov_gt_tr_only}": cov_gt_tr_only, "{cov_gt_new}":     cov_gt_new,
    "{cov_gt_loaded}":  cov_gt_loaded,  "{cov_gt_total}":   cov_gt_total,   "{cov_gt_rows}":    cov_gt_rows,
    "{cov_gt_target}":  cov_gt_target,  "{cov_gt_pct}":     cov_gt_pct,
    "{cov_gc_bl}":      cov_gc_bl,      "{cov_gc_new}":     cov_gc_new,
    "{cov_gc_loaded}":  cov_gc_loaded,  "{cov_gc_total}":   cov_gc_total,   "{cov_gc_rows}":    cov_gc_rows,
    "{cov_gc_target}":  cov_gc_target,  "{cov_gc_pct}":     cov_gc_pct,
    "{cov_bt_trained}": cov_bt_trained, "{cov_bt_new}":     cov_bt_new,
    "{cov_bt_loaded}":  cov_bt_loaded,  "{cov_bt_total}":   cov_bt_total,   "{cov_bt_rows}":    cov_bt_rows,
    "{cov_bt_target}":  cov_bt_target,  "{cov_bt_pct}":     cov_bt_pct,
    "{cov_bt_status}":  ("Awaiting first boys roster file." if not cov_bt_has_rosters else
                        f"{cov_bt_loaded} of {cov_bt_total} schools have rosters loaded."),
    "{cov_gt_school_rows}": cov_gt_school_rows,
    "{cov_gc_school_rows}": cov_gc_school_rows,
    "{cov_bt_school_rows}": cov_bt_school_rows,
}
for k, v in subs.items():
    output_text = output_text.replace(k, str(v))

# Update last-updated timestamp in the footer line
output_text = output_text.replace(
    "Data collected April 2026 &nbsp;·&nbsp; Pakistan Urban Schools",
    f"Data updated {last_updated} &nbsp;·&nbsp; {n_total} records ({n_boys} boys · {n_girls} girls) &nbsp;·&nbsp; Pakistan Urban Schools"
)

OUTPUT.write_text(output_text, encoding="utf-8")
print(f"index.html written ({OUTPUT.stat().st_size // 1024} KB)")

# ══════════════════════════════════════════════════════════════════════════
# GIT COMMIT & PUSH
# ══════════════════════════════════════════════════════════════════════════
if AUTO_PUSH:
    print("Committing and pushing to GitHub...")
    cmds = [
        ["git", "add", "index.html"],
        ["git", "commit", "-m", f"Dashboard update {last_updated} — {n_total} records ({n_boys}B/{n_girls}G)"],
        ["git", "push"],
    ]
    for cmd in cmds:
        result = subprocess.run(cmd, cwd=BASE, capture_output=True, text=True)
        if result.returncode != 0:
            print(f"  WARNING: {' '.join(cmd)} failed:\n  {result.stderr.strip()}")
        else:
            print(f"  OK: {' '.join(cmd)}")
    print(f"\nDone! Dashboard live at: https://uhs.rs.org.pk")
else:
    print("\nDone (local only — skipped git push)")
    print(f"Open D:/UHS Project/index.html to preview")
